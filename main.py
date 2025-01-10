from duckduckgo_search import DDGS
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement


data = ["Зарубежная Азия", "Африка", "Северная Америка (США, Канада)", "Латинская Америка", "Австралия и Океания"]  # темы

addon = {
    "Состав и государственный строй": ["Состав, Страны монархии, Страны республики. Пиши каждый пункт в одну строку"],  # подтема и как писать
    "Природные ресурсы": ["все природные ресурсы"],
    "Население": ["население (общее население, состав, ЕП, уровень жизни и образования)"],
    "Хозяйство": ["хозяйство"]
}


for dda in data:  # Новый документ dda - название и главный заголовок
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    heading = doc.add_paragraph(dda)  # Добавляем главный заголовок документа
    run = heading.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

    for vo in addon:
        heading = doc.add_paragraph(f"{vo}: ")  # Добавляем главный заголовок документа
        run = heading.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True

        for te in addon[vo]:
            if len(addon[vo]) > 1:
                heading = doc.add_paragraph(f"{te}: ")  # Добавляем главный заголовок документа
                run = heading.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True
            text = DDGS().chat(f'{dda}. Напиши: {te}{". Пиши без вывода" if len(addon[vo]) > 1 else ""}', model='claude-3-haiku')
            heading = doc.add_paragraph(text+" ")  # Добавляем текст
            run = heading.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    doc.save(f'{dda}.docx')
